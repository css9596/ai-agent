"""다양한 형식으로 분석 결과 내보내기"""

import json
import base64
from pathlib import Path
from datetime import datetime
from typing import Optional

import markdown2
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# WeasyPrint은 시스템 의존성이 많아서 선택적으로 로드
HAS_WEASYPRINT = False
try:
    import weasyprint
    HAS_WEASYPRINT = True
except (ImportError, OSError):
    pass


def export_to_html(
    markdown_content: str,
    title: str = "개발 분석서",
    author: str = "AI Assistant"
) -> str:
    """마크다운을 HTML로 변환 (스타일 포함)"""

    # 마크다운을 HTML로 변환
    html_content = markdown2.markdown(markdown_content, extras=['tables', 'fenced-code-blocks'])

    # CSS 스타일
    css = """
        <style>
            @import url('https://fonts.googleapis.com/css2?family=Noto+Sans+KR:wght@400;600;700&display=swap');

            * {
                margin: 0;
                padding: 0;
                box-sizing: border-box;
            }

            html, body {
                font-family: 'Noto Sans KR', 'Noto Sans CJK KR', 'Malgun Gothic', 'Apple SD Gothic Neo', -apple-system, BlinkMacSystemFont, sans-serif;
                line-height: 1.7;
                color: #1f2937;
            }

            body {
                max-width: 900px;
                margin: 0 auto;
                padding: 40px 20px;
                background: #f9fafb;
            }

            .header {
                background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
                color: white;
                padding: 40px;
                border-radius: 12px;
                margin-bottom: 40px;
                text-align: center;
            }

            .header h1 {
                font-size: 32px;
                margin-bottom: 10px;
            }

            .metadata {
                font-size: 13px;
                opacity: 0.9;
                margin-top: 15px;
            }

            .content {
                background: white;
                padding: 40px;
                border-radius: 12px;
                box-shadow: 0 1px 3px rgba(0,0,0,0.1);
            }

            h1 {
                font-size: 28px;
                margin: 30px 0 20px;
                border-bottom: 3px solid #667eea;
                padding-bottom: 10px;
                color: #1f2937;
            }

            h2 {
                font-size: 22px;
                margin: 25px 0 15px;
                color: #374151;
                border-left: 4px solid #667eea;
                padding-left: 12px;
            }

            h3 {
                font-size: 18px;
                margin: 20px 0 10px;
                color: #4b5563;
            }

            p {
                margin: 12px 0;
                text-align: justify;
            }

            ul, ol {
                margin: 12px 0 12px 30px;
            }

            li {
                margin: 6px 0;
            }

            code {
                background: #f3f4f6;
                padding: 2px 6px;
                border-radius: 4px;
                font-family: 'Courier New', monospace;
                font-size: 13px;
                color: #c7254e;
            }

            pre {
                background: #1e293b;
                color: #e2e8f0;
                padding: 16px;
                border-radius: 8px;
                overflow-x: auto;
                margin: 12px 0;
            }

            pre code {
                background: none;
                color: inherit;
                padding: 0;
            }

            table {
                width: 100%;
                border-collapse: collapse;
                margin: 16px 0;
                font-size: 14px;
                border: 1px solid #e5e7eb;
            }

            th {
                background: #f3f4f6;
                padding: 12px;
                text-align: left;
                font-weight: 600;
                border: 1px solid #e5e7eb;
            }

            td {
                padding: 12px;
                border: 1px solid #e5e7eb;
            }

            tr:hover {
                background: #f9fafb;
            }

            blockquote {
                border-left: 4px solid #667eea;
                padding-left: 16px;
                margin: 12px 0;
                color: #6b7280;
                font-style: italic;
            }

            a {
                color: #667eea;
                text-decoration: none;
                border-bottom: 1px solid #e5e7eb;
            }

            a:hover {
                color: #764ba2;
            }

            .page-break {
                page-break-after: always;
            }

            @media print {
                body {
                    background: white;
                }

                .content {
                    box-shadow: none;
                    border: none;
                }
            }
        </style>
    """

    # 전체 HTML 작성
    html = f"""<!DOCTYPE html>
<html lang="ko">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title}</title>
    {css}
</head>
<body>
    <div class="header">
        <h1>{title}</h1>
        <div class="metadata">
            <p>작성자: {author}</p>
            <p>생성일시: {datetime.now().strftime('%Y년 %m월 %d일 %H:%M:%S')}</p>
        </div>
    </div>

    <div class="content">
        {html_content}
    </div>
</body>
</html>"""

    return html


def export_to_pdf(markdown_content: str, output_path: str) -> bool:
    """마크다운을 PDF로 변환"""
    if not HAS_WEASYPRINT:
        print("PDF 변환 실패: WeasyPrint가 설치되지 않았습니다. 시스템 의존성을 확인하세요.")
        return False

    try:
        html = export_to_html(markdown_content)
        weasyprint.HTML(string=html).write_pdf(output_path)
        return True
    except Exception as e:
        print(f"PDF 변환 실패: {e}")
        return False


def export_to_docx(markdown_content: str, output_path: str) -> bool:
    """마크다운을 Word (.docx)로 변환"""
    try:
        doc = Document()

        # 제목
        title = doc.add_heading("개발 분석서", level=1)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # 메타데이터
        metadata = doc.add_paragraph()
        metadata.add_run(f"작성자: AI Assistant\n")
        metadata.add_run(f"생성일시: {datetime.now().strftime('%Y년 %m월 %d일 %H:%M:%S')}\n")
        metadata.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # 페이지 나누기
        doc.add_page_break()

        # 마크다운 파싱
        lines = markdown_content.split("\n")
        i = 0
        while i < len(lines):
            line = lines[i]

            # 헤딩
            if line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            # 테이블
            elif line.startswith("|"):
                table_lines = []
                while i < len(lines) and lines[i].startswith("|"):
                    table_lines.append(lines[i])
                    i += 1
                i -= 1  # 다음 반복을 위해 인덱스 조정

                if table_lines:
                    # 테이블 파싱
                    rows = [row.split("|")[1:-1] for row in table_lines]
                    if len(rows) > 0:
                        table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                        table.style = "Light Grid Accent 1"

                        for row_idx, row_data in enumerate(rows):
                            for col_idx, cell_data in enumerate(row_data):
                                cell = table.rows[row_idx].cells[col_idx]
                                cell.text = cell_data.strip()

            # 코드 블록
            elif line.startswith("```"):
                code_lines = []
                i += 1
                while i < len(lines) and not lines[i].startswith("```"):
                    code_lines.append(lines[i])
                    i += 1

                code_block = doc.add_paragraph()
                code_text = code_block.add_run("\n".join(code_lines))
                code_text.font.name = "Courier New"
                code_text.font.size = Pt(10)
                code_block.paragraph_format.left_indent = Inches(0.5)

            # 목록
            elif line.startswith("- "):
                doc.add_paragraph(line[2:], style="List Bullet")
            elif line.startswith("* "):
                doc.add_paragraph(line[2:], style="List Bullet")

            # 일반 텍스트
            elif line.strip():
                doc.add_paragraph(line)

            i += 1

        # 저장
        doc.save(output_path)
        return True

    except Exception as e:
        print(f"Word 변환 실패: {e}")
        return False


def export_to_json(data: dict, output_path: str) -> bool:
    """분석 결과를 JSON으로 저장"""
    try:
        with open(output_path, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        return True
    except Exception as e:
        print(f"JSON 저장 실패: {e}")
        return False


def get_export_filename(base_name: str, format: str) -> str:
    """내보내기 파일명 생성"""
    ext_map = {
        "html": "html",
        "pdf": "pdf",
        "docx": "docx",
        "json": "json",
    }

    ext = ext_map.get(format, format)
    name_without_ext = base_name.replace(".md", "")
    return f"{name_without_ext}.{ext}"
